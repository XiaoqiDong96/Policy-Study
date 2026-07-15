#!/usr/bin/env python3
"""Build auditable prefecture-level variables from downloaded official lists.

Scope
-----
* Ministry of Education annual higher-school lists.
* Ministry of Science and Technology national incubator lists.
* Ministry of Culture and Tourism intangible-cultural-heritage and cultural-
  industry lists.
* MIIT culture/technology-fusion and future-industry attachments.

Air quality is deliberately excluded.  The output universe is the complete
mainland prefecture-level administrative universe (333 prefecture-level units)
plus the four municipalities as prefecture-equivalent observations: 337 rows
per year.  The universe is parsed from the Ministry of Civil Affairs (MCA)
National Administrative Division Information Query Platform.

No fuzzy location cleaning, imputation, or silent deduplication is performed.
Exact official names, deterministic suffix-free place tokens, and current MCA
county-to-prefecture code relationships are the only automatic mappings.  Any
suffix-free token match is retained with a flag.  Unsafe zeros remain blank.

Runtime requirements
--------------------
Python packages: lxml, openpyxl, pdfplumber, python-docx.
System command: LibreOffice/soffice (for legacy .xls and .doc conversion).

The bundled Codex Python runtime contains the Python packages on the author's
workstation.  Example:

    /path/to/python build_official_lists_panel.py --project-root /path/project
"""

from __future__ import annotations

import argparse
import csv
import datetime as dt
import hashlib
import html as html_lib
import json
import re
import shutil
import subprocess
import sys
import tempfile
import urllib.request
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable, Iterator, Sequence

try:
    import openpyxl
    import pdfplumber
    from docx import Document
    from lxml import html as lxml_html
except ImportError as exc:  # pragma: no cover - environment diagnostic
    raise SystemExit(
        "Missing parser dependency. Required: lxml, openpyxl, pdfplumber, "
        f"python-docx. Original error: {exc}"
    ) from exc


MCA_URL = "http://xzqh.mca.gov.cn/map"
MUNICIPALITY_CODES = {"110000", "120000", "310000", "500000"}
EXCLUDED_PROVINCE_PREFIXES = {"71", "81", "82"}
VARIABLE_SPECS = {
    "moe_general_university_count": {
        "label": "Ordinary higher-education institutions (annual stock)",
        "index_eligible": True,
        "construct": "knowledge and talent connection infrastructure",
        "status": "confirmed",
    },
    "moe_adult_higher_school_count": {
        "label": "Adult higher-education institutions (annual stock)",
        "index_eligible": False,
        "construct": "lifelong-learning infrastructure; sensitivity only",
        "status": "confirmed_but_location_incomplete",
    },
    "most_new_national_incubator_count": {
        "label": "Newly announced national technology incubators (annual flow)",
        "index_eligible": True,
        "construct": "entrepreneurial ecosystem and incubation",
        "status": "confirmed",
    },
    "mct_active_intangible_protection_unit_count": {
        "label": "Active national intangible-heritage protection units (stock)",
        "index_eligible": True,
        "construct": "cultural continuity and historical development",
        "status": "confirmed",
    },
    "mct_intangible_productive_base_count": {
        "label": "National productive-protection demonstration bases",
        "index_eligible": False,
        "construct": "cultural industry and living heritage; sensitivity only",
        "status": "confirmed_but_city_location_incomplete",
    },
    "mct_cultural_industry_proposed_new_base_count": {
        "label": "Proposed new national cultural-industry demonstration bases",
        "index_eligible": False,
        "construct": "cultural industry; proposal-stage candidate only",
        "status": "proposal_not_final",
    },
    "mct_cultural_industry_proposed_retained_base_count": {
        "label": "Proposed retained national cultural-industry demonstration bases",
        "index_eligible": False,
        "construct": "cultural industry; proposal-stage candidate only",
        "status": "proposal_not_final",
    },
    "miit_culture_technology_fusion_base_count": {
        "label": "National culture/technology-integration demonstration bases",
        "index_eligible": True,
        "construct": "culture-technology integration",
        "status": "confirmed",
    },
}

RECORD_FIELDS = [
    "record_id",
    "source_id",
    "list_year",
    "data_period",
    "variable",
    "list_name",
    "list_status",
    "category",
    "province_raw",
    "location_raw",
    "entity_name",
    "item_name",
    "detail_type",
    "aggregation_key",
    "source_file",
    "source_url",
    "source_sha256",
    "extraction_method",
    "mapping_text",
    "prefecture_code",
    "prefecture_name",
    "prefecture_type",
    "province_code",
    "province_name",
    "match_status",
    "match_method",
    "match_evidence",
    "usable_for_city_panel",
    "index_eligible",
    "qc_flags",
]


def clean_text(value: Any) -> str:
    """Whitespace-only normalization; content is otherwise unchanged."""
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def stable_record_id(parts: Sequence[Any]) -> str:
    raw = "\x1f".join(clean_text(part) for part in parts).encode("utf-8")
    return hashlib.sha256(raw).hexdigest()[:20]


def write_csv(path: Path, rows: Iterable[dict[str, Any]], fields: Sequence[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(fields), extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def read_download_registry(project_root: Path) -> dict[str, dict[str, str]]:
    path = project_root / "01_source_register" / "download_events.csv"
    registry: dict[str, dict[str, str]] = {}
    with path.open(encoding="utf-8-sig", newline="") as handle:
        for row in csv.DictReader(handle):
            saved = row.get("saved_path", "")
            if saved:
                registry[saved] = row
    return registry


def source_metadata(project_root: Path, path: Path, registry: dict[str, dict[str, str]]) -> dict[str, str]:
    rel = path.relative_to(project_root).as_posix()
    row = registry.get(rel, {})
    return {
        "source_file": rel,
        "source_url": row.get("url", ""),
        "source_sha256": row.get("sha256", "") or sha256_file(path),
    }


def locate_soffice() -> str:
    for command in ("soffice", "libreoffice"):
        found = shutil.which(command)
        if found:
            return found
    raise SystemExit("LibreOffice/soffice is required to read legacy .xls and .doc files.")


class LegacyConverter:
    """Hash-keyed temporary conversion cache for legacy Office documents."""

    def __init__(self, cache_dir: Path):
        self.cache_dir = cache_dir
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        self.soffice = locate_soffice()

    def convert(self, source: Path, target_ext: str) -> Path:
        digest = sha256_file(source)[:16]
        target = self.cache_dir / f"{source.stem}_{digest}.{target_ext}"
        if target.exists() and target.stat().st_size:
            return target
        with tempfile.TemporaryDirectory(prefix="official_lists_convert_") as tmp:
            tmp_dir = Path(tmp)
            command = [
                self.soffice,
                "--headless",
                "--convert-to",
                target_ext,
                "--outdir",
                str(tmp_dir),
                str(source),
            ]
            completed = subprocess.run(command, capture_output=True, text=True, check=False)
            candidates = list(tmp_dir.glob(f"*.{target_ext}"))
            if completed.returncode != 0 or len(candidates) != 1:
                raise RuntimeError(
                    f"Conversion failed for {source}: return={completed.returncode}; "
                    f"stdout={completed.stdout[-500:]}; stderr={completed.stderr[-500:]}"
                )
            shutil.copy2(candidates[0], target)
        return target


def fetch_mca_snapshot(snapshot_path: Path, refresh: bool) -> tuple[str, dict[str, str]]:
    """Return MCA HTML decoded as GB18030 and its provenance metadata."""
    snapshot_path.parent.mkdir(parents=True, exist_ok=True)
    if refresh or not snapshot_path.exists():
        request = urllib.request.Request(
            MCA_URL,
            headers={"User-Agent": "Mozilla/5.0 official-list-research/1.0"},
        )
        with urllib.request.urlopen(request, timeout=60) as response:
            payload = response.read()
        snapshot_path.write_bytes(payload)
        fetched_at = dt.datetime.now(dt.timezone.utc).isoformat()
    else:
        payload = snapshot_path.read_bytes()
        fetched_at = dt.datetime.fromtimestamp(
            snapshot_path.stat().st_mtime, tz=dt.timezone.utc
        ).isoformat()
    text = payload.decode("gb18030", errors="replace")
    meta = {
        "source_url": MCA_URL,
        "snapshot_file": snapshot_path.as_posix(),
        "snapshot_sha256": hashlib.sha256(payload).hexdigest(),
        "snapshot_fetched_or_mtime_utc": fetched_at,
        "effective_date": "[NOT EXPOSED BY SOURCE PAGE]",
        "source_note": (
            "MCA page states that content is compiled from locally published information "
            "and updated as those publications change."
        ),
    }
    return text, meta


def parse_mca_divisions(page_text: str) -> list[dict[str, str]]:
    match = re.search(r"value='(\[\{.*?\}\])'\s+id=\"pyArr\"", page_text, re.S)
    if not match:
        raise ValueError("Could not find MCA pyArr administrative-division payload.")
    payload = html_lib.unescape(match.group(1))
    divisions = json.loads(payload)
    if len(divisions) < 3000:
        raise ValueError(f"MCA division payload is unexpectedly short: {len(divisions)}")
    required = {"cName", "code"}
    if any(not required.issubset(row) for row in divisions):
        raise ValueError("MCA division payload lacks cName/code fields.")
    return [{key: clean_text(value) for key, value in row.items()} for row in divisions]


def classify_prefecture(name: str, code: str) -> str:
    if code in MUNICIPALITY_CODES:
        return "municipality_prefecture_equivalent"
    if name.endswith("自治州"):
        return "autonomous_prefecture"
    if name.endswith("地区"):
        return "prefecture"
    if name.endswith("盟"):
        return "league"
    if name.endswith("市"):
        return "prefecture_level_city"
    return "[NEEDS REVIEW]"


def build_prefecture_universe(divisions: list[dict[str, str]]) -> list[dict[str, str]]:
    by_code = {row["code"]: row for row in divisions}
    province_by_prefix = {
        row["code"][:2]: row
        for row in divisions
        if row["code"].endswith("0000")
        and row["code"][:2] not in EXCLUDED_PROVINCE_PREFIXES
    }
    prefectures = [
        row
        for row in divisions
        if row["code"].endswith("00")
        and not row["code"].endswith("0000")
        and row["code"][:2] not in EXCLUDED_PROVINCE_PREFIXES
    ]
    prefectures.extend(by_code[code] for code in sorted(MUNICIPALITY_CODES))
    result = []
    for row in sorted(prefectures, key=lambda item: item["code"]):
        province = province_by_prefix[row["code"][:2]]
        result.append(
            {
                "prefecture_code": row["code"],
                "prefecture_name": row["cName"],
                "prefecture_type": classify_prefecture(row["cName"], row["code"]),
                "province_code": province["code"],
                "province_name": province["cName"],
                "pinyin": row.get("py", ""),
                "universe_scope": "mainland_333_prefectures_plus_4_municipalities",
            }
        )
    if len(result) != 337:
        raise ValueError(f"Expected 337 prefecture-equivalent units; found {len(result)}")
    return result


def province_aliases(name: str) -> set[str]:
    aliases = {name}
    for suffix in ("壮族自治区", "维吾尔自治区", "回族自治区", "自治区", "省", "市"):
        if name.endswith(suffix):
            aliases.add(name[: -len(suffix)])
    return {alias for alias in aliases if alias}


@dataclass(frozen=True)
class MatchResult:
    prefecture_code: str = ""
    prefecture_name: str = ""
    prefecture_type: str = ""
    province_code: str = ""
    province_name: str = ""
    status: str = "unmatched"
    method: str = ""
    evidence: str = ""
    flags: str = ""


class DivisionMatcher:
    def __init__(
        self,
        divisions: list[dict[str, str]],
        universe: list[dict[str, str]],
    ) -> None:
        self.universe = universe
        self.pref_by_code = {row["prefecture_code"]: row for row in universe}
        self.province_by_alias: dict[str, tuple[str, str]] = {}
        provinces = [
            row
            for row in divisions
            if row["code"].endswith("0000")
            and row["code"][:2] not in EXCLUDED_PROVINCE_PREFIXES
        ]
        for province in provinces:
            for alias in province_aliases(province["cName"]):
                self.province_by_alias[alias] = (province["code"], province["cName"])

        self.pref_full_names = sorted(
            ((row["prefecture_name"], row["prefecture_code"]) for row in universe),
            key=lambda item: (-len(item[0]), item[1]),
        )
        self.pref_aliases: list[tuple[str, str]] = []
        self.pref_prefix_aliases: list[tuple[str, str]] = []
        for row in universe:
            name = row["prefecture_name"]
            aliases = {name}
            for suffix in ("自治州", "地区", "盟", "市"):
                if name.endswith(suffix):
                    aliases.add(name[: -len(suffix)])
            for alias in aliases:
                if alias != name and len(alias) >= 2:
                    self.pref_prefix_aliases.append((alias, row["prefecture_code"]))
                # Substring matching is more restrictive than start-of-field
                # matching to avoid common-word false positives such as "中山".
                if len(alias) >= 3 or row["prefecture_code"] in MUNICIPALITY_CODES:
                    if alias != name:
                        self.pref_aliases.append((alias, row["prefecture_code"]))
        self.pref_aliases.sort(key=lambda item: (-len(item[0]), item[1]))
        self.pref_prefix_aliases.sort(key=lambda item: (-len(item[0]), item[1]))

        self.counties: list[tuple[str, str, str]] = []
        self.direct_counties: list[tuple[str, str, str, str, str]] = []
        for division in divisions:
            code = division["code"]
            if code.endswith("00") or code[:2] in EXCLUDED_PROVINCE_PREFIXES:
                continue
            if code[:2] in {"11", "12", "31", "50"}:
                parent_code = code[:2] + "0000"
            else:
                parent_code = code[:4] + "00"
            if parent_code in self.pref_by_code:
                self.counties.append((division["cName"], parent_code, code[:2]))
            else:
                province = next(
                    (
                        row
                        for row in universe
                        if row["province_code"].startswith(code[:2])
                    ),
                    None,
                )
                if province:
                    self.direct_counties.append(
                        (
                            division["cName"],
                            division["code"],
                            province["province_code"],
                            province["province_name"],
                            code[:2],
                        )
                    )
        self.counties.sort(key=lambda item: (-len(item[0]), item[1], item[0]))
        self.direct_counties.sort(key=lambda item: (-len(item[0]), item[1]))

    def find_province(self, text: str) -> tuple[str, str]:
        text = clean_text(text)
        exact = self.province_by_alias.get(text)
        if exact:
            return exact
        candidates = [
            value
            for alias, value in self.province_by_alias.items()
            if len(alias) >= 2 and alias in text
        ]
        unique = sorted(set(candidates))
        return unique[0] if len(unique) == 1 else ("", "")

    def _result(self, code: str, status: str, method: str, evidence: str, flags: str = "") -> MatchResult:
        row = self.pref_by_code[code]
        return MatchResult(
            prefecture_code=row["prefecture_code"],
            prefecture_name=row["prefecture_name"],
            prefecture_type=row["prefecture_type"],
            province_code=row["province_code"],
            province_name=row["province_name"],
            status=status,
            method=method,
            evidence=evidence,
            flags=flags,
        )

    def _unique_candidates(self, candidates: Iterable[tuple[str, str]]) -> list[tuple[str, str]]:
        by_code: dict[str, str] = {}
        for evidence, code in candidates:
            by_code.setdefault(code, evidence)
        return [(evidence, code) for code, evidence in sorted(by_code.items())]

    def _longest_candidates(self, candidates: Iterable[tuple[str, str]]) -> list[tuple[str, str]]:
        unique = self._unique_candidates(candidates)
        if not unique:
            return []
        longest = max(len(evidence) for evidence, _ in unique)
        return [item for item in unique if len(item[0]) == longest]

    def match(self, province_raw: str, location_text: str, fallback_text: str) -> MatchResult:
        province_code, province_name = self.find_province(province_raw or location_text)
        province_prefix = province_code[:2] if province_code else ""
        location_text = clean_text(location_text)
        fallback_text = clean_text(fallback_text)

        exact_location = self._longest_candidates(
            (name, code) for name, code in self.pref_full_names if name in location_text
        )
        if len(exact_location) == 1:
            evidence, code = exact_location[0]
            return self._result(code, "matched_exact", "official_prefecture_name_in_location", evidence)
        if len(exact_location) > 1:
            return MatchResult(
                province_code=province_code,
                province_name=province_name,
                status="ambiguous",
                method="multiple_official_prefecture_names_in_location",
                evidence="|".join(item[0] for item in exact_location),
            )

        # Some national lists put separately planned cities in a province-like
        # column (e.g. 大连, 宁波, 厦门, 青岛, 深圳).  Exact cell equality only.
        exact_alias_cell = self._unique_candidates(
            (alias, code) for alias, code in self.pref_aliases if province_raw == alias
        )
        if len(exact_alias_cell) == 1:
            evidence, code = exact_alias_cell[0]
            return self._result(
                code,
                "matched_flagged",
                "exact_suffix_free_prefecture_cell",
                evidence,
                "suffix_free_place_token",
            )

        county_candidates = []
        for county_name, parent_code, county_province in self.counties:
            if province_prefix and county_province != province_prefix:
                continue
            if county_name in location_text:
                county_candidates.append((county_name, parent_code))
        unique_counties = self._unique_candidates(county_candidates)
        if len(unique_counties) == 1:
            evidence, code = unique_counties[0]
            return self._result(code, "matched_exact", "official_county_code_parent", evidence)
        if len(unique_counties) > 1:
            return MatchResult(
                province_code=province_code,
                province_name=province_name,
                status="ambiguous",
                method="multiple_county_parent_candidates",
                evidence="|".join(item[0] for item in unique_counties),
            )

        direct_candidates = [
            (name, code, pcode, pname)
            for name, code, pcode, pname, prefix in self.direct_counties
            if (not province_prefix or prefix == province_prefix)
            and (name in location_text or name in fallback_text)
        ]
        direct_codes = {item[1] for item in direct_candidates}
        if len(direct_codes) == 1:
            evidence, code, pcode, pname = direct_candidates[0]
            return MatchResult(
                province_code=pcode,
                province_name=pname,
                status="outside_prefecture_universe",
                method="official_province_direct_county_level_unit",
                evidence=f"{evidence}[{code}]",
                flags="outside_prefecture_level_scope",
            )

        exact_fallback = self._longest_candidates(
            (name, code) for name, code in self.pref_full_names if name in fallback_text
        )
        if len(exact_fallback) == 1:
            evidence, code = exact_fallback[0]
            return self._result(
                code,
                "matched_exact",
                "official_prefecture_name_in_official_item_text",
                evidence,
            )
        if len(exact_fallback) > 1:
            return MatchResult(
                province_code=province_code,
                province_name=province_name,
                status="ambiguous",
                method="multiple_prefecture_names_in_official_item_text",
                evidence="|".join(item[0] for item in exact_fallback),
            )

        prefix_alias_fallback = []
        for alias, code in self.pref_prefix_aliases:
            if province_prefix and not code.startswith(province_prefix):
                continue
            if re.search(rf"(?:^|\s|[（(]){re.escape(alias)}", fallback_text):
                prefix_alias_fallback.append((alias, code))
        unique_prefix_aliases = self._longest_candidates(prefix_alias_fallback)
        if len(unique_prefix_aliases) == 1:
            evidence, code = unique_prefix_aliases[0]
            return self._result(
                code,
                "matched_flagged",
                "suffix_free_prefecture_token_at_official_field_start",
                evidence,
                "suffix_free_place_token",
            )
        if len(unique_prefix_aliases) > 1:
            return MatchResult(
                province_code=province_code,
                province_name=province_name,
                status="ambiguous",
                method="multiple_suffix_free_prefecture_field_prefixes",
                evidence="|".join(item[0] for item in unique_prefix_aliases),
                flags="suffix_free_place_token",
            )

        alias_fallback = []
        for alias, code in self.pref_aliases:
            if province_prefix and not code.startswith(province_prefix):
                continue
            if alias in fallback_text:
                alias_fallback.append((alias, code))
        unique_aliases = self._longest_candidates(alias_fallback)
        if len(unique_aliases) == 1:
            evidence, code = unique_aliases[0]
            return self._result(
                code,
                "matched_flagged",
                "suffix_free_prefecture_token_in_official_item_text",
                evidence,
                "suffix_free_place_token",
            )
        if len(unique_aliases) > 1:
            return MatchResult(
                province_code=province_code,
                province_name=province_name,
                status="ambiguous",
                method="multiple_suffix_free_prefecture_tokens",
                evidence="|".join(item[0] for item in unique_aliases),
                flags="suffix_free_place_token",
            )

        fallback_counties = []
        for county_name, parent_code, county_province in self.counties:
            if province_prefix and county_province != province_prefix:
                continue
            if county_name in fallback_text:
                fallback_counties.append((county_name, parent_code))
        unique_fallback_counties = self._unique_candidates(fallback_counties)
        if len(unique_fallback_counties) == 1:
            evidence, code = unique_fallback_counties[0]
            return self._result(
                code,
                "matched_flagged",
                "official_county_name_in_official_item_text",
                evidence,
                "county_derived_from_entity_text",
            )

        if province_code:
            return MatchResult(
                province_code=province_code,
                province_name=province_name,
                status="unmatched_province_only",
                method="province_only",
                evidence=province_raw,
            )
        return MatchResult(status="unmatched", method="no_official_place_match")


def iter_xlsx_rows(path: Path, max_col: int = 10) -> Iterator[list[str]]:
    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    worksheet = workbook.active
    for row in worksheet.iter_rows(max_col=max_col, values_only=True):
        yield [clean_text(value) for value in row]


def base_record(
    *,
    source_id: str,
    list_year: int,
    data_period: str,
    variable: str,
    list_name: str,
    list_status: str,
    category: str,
    province_raw: str,
    location_raw: str,
    entity_name: str,
    item_name: str,
    detail_type: str,
    aggregation_key: str,
    source: dict[str, str],
    extraction_method: str,
    mapping_text: str,
) -> dict[str, Any]:
    record = {
        "source_id": source_id,
        "list_year": list_year,
        "data_period": data_period,
        "variable": variable,
        "list_name": list_name,
        "list_status": list_status,
        "category": category,
        "province_raw": province_raw,
        "location_raw": location_raw,
        "entity_name": entity_name,
        "item_name": item_name,
        "detail_type": detail_type,
        "aggregation_key": aggregation_key or entity_name or item_name,
        **source,
        "extraction_method": extraction_method,
        "mapping_text": mapping_text,
    }
    record["record_id"] = stable_record_id(
        [source_id, list_year, list_status, aggregation_key, entity_name, item_name, source["source_file"]]
    )
    return record


def extract_moe(
    project_root: Path,
    converter: LegacyConverter,
    registry: dict[str, dict[str, str]],
) -> list[dict[str, Any]]:
    records = []
    for source_path in sorted((project_root / "03_external_raw" / "moe_universities").glob("*.xls")):
        converted = converter.convert(source_path, "xlsx")
        rows = list(iter_xlsx_rows(converted, max_col=8))
        title = " ".join(cell for row in rows[:4] for cell in row if cell)
        year_match = re.search(r"截至\s*(20\d{2})年", title)
        source = source_metadata(project_root, source_path, registry)
        if not year_match:
            # The 2019 attachments omit the date inside the workbook, but the
            # official parent-page URL is explicitly under /201906/.  This is
            # source metadata, not a filename/content guess.
            event = registry.get(source["source_file"], {})
            year_match = re.search(
                r"/(20\d{2})(?:\d{2})?/",
                event.get("parent_url", "") or event.get("url", ""),
            )
        if not year_match:
            continue
        year = int(year_match.group(1))
        is_adult = "成人高等学校名单" in title
        is_general = "普通高等学校名单" in title
        if not (is_adult or is_general):
            continue
        variable = "moe_adult_higher_school_count" if is_adult else "moe_general_university_count"
        list_name = "全国成人高等学校名单" if is_adult else "全国普通高等学校名单"
        header_index = next(
            (index for index, row in enumerate(rows) if "学校名称" in row and "学校标识码" in row),
            None,
        )
        if header_index is None:
            continue
        header = rows[header_index]
        columns = {name: header.index(name) for name in header if name}
        current_province = ""
        for row in rows[header_index + 1 :]:
            first = row[0]
            section = re.match(r"^(.+?)(?:（|\()\s*\d+\s*所", first)
            if section:
                current_province = section.group(1)
                continue
            seq = row[columns.get("序号", 0)] if row else ""
            if not re.fullmatch(r"\d+", seq):
                continue
            entity = row[columns["学校名称"]]
            school_id = row[columns["学校标识码"]]
            location = row[columns["所在地"]] if "所在地" in columns else current_province
            level = row[columns["办学层次"]] if "办学层次" in columns else ""
            supervisor = row[columns["主管部门"]] if "主管部门" in columns else ""
            records.append(
                base_record(
                    source_id="moe_universities",
                    list_year=year,
                    data_period=str(year),
                    variable=variable,
                    list_name=list_name,
                    list_status="confirmed_stock",
                    category="knowledge_connection",
                    province_raw=current_province,
                    location_raw=location,
                    entity_name=entity,
                    item_name=school_id,
                    detail_type=f"level={level}; supervisor={supervisor}",
                    aggregation_key=school_id or entity,
                    source=source,
                    extraction_method="soffice_xls_to_xlsx_openpyxl",
                    mapping_text=" ".join([entity, location, current_province]),
                )
            )
    return records


def html_table_rows(path: Path) -> list[list[str]]:
    document = lxml_html.fromstring(path.read_bytes())
    tables = []
    for table in document.xpath("//table"):
        rows = []
        for tr in table.xpath(".//tr"):
            cells = [clean_text(cell.text_content()) for cell in tr.xpath("./th|./td")]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)
    return max(tables, key=len) if tables else []


def extract_incubator_html(
    project_root: Path,
    registry: dict[str, dict[str, str]],
) -> list[dict[str, Any]]:
    records = []
    for year in (2015, 2016, 2017):
        directory = project_root / "03_external_raw" / f"most_incubator_{year}"
        for path in sorted(directory.glob("*.html")):
            rows = html_table_rows(path)
            if not rows or "序号" not in "".join(rows[0]):
                continue
            current_region = ""
            source = source_metadata(project_root, path, registry)
            for row in rows[1:]:
                if not row or not re.fullmatch(r"\d+", row[0]):
                    continue
                if len(row) >= 3:
                    current_region = row[1] or current_region
                    entity = row[2]
                elif len(row) == 2:
                    entity = row[1]
                else:
                    continue
                records.append(
                    base_record(
                        source_id=f"most_incubator_{year}",
                        list_year=year,
                        data_period=str(year),
                        variable="most_new_national_incubator_count",
                        list_name=f"{year}年度国家级科技企业孵化器名单",
                        list_status="confirmed_flow",
                        category="entrepreneurial_ecosystem",
                        province_raw=current_region,
                        location_raw=current_region,
                        entity_name=entity,
                        item_name="",
                        detail_type="",
                        aggregation_key=entity,
                        source=source,
                        extraction_method="lxml_html_table_with_merged_cell_carry_forward",
                        mapping_text=" ".join([current_region, entity]),
                    )
                )
    return records


def docx_table(path: Path) -> list[list[str]]:
    document = Document(path)
    if not document.tables:
        return []
    table = max(document.tables, key=lambda item: len(item.rows))
    return [[clean_text(cell.text) for cell in row.cells] for row in table.rows]


def extract_incubator_doc(
    project_root: Path,
    converter: LegacyConverter,
    registry: dict[str, dict[str, str]],
) -> list[dict[str, Any]]:
    records = []
    for year in (2020, 2022):
        directory = project_root / "03_external_raw" / f"most_incubator_{year}"
        for source_path in sorted(directory.glob("*.doc")):
            converted = converter.convert(source_path, "docx")
            rows = docx_table(converted)
            if not rows:
                continue
            header = rows[0]
            if "孵化器名称" not in header:
                continue
            columns = {name: header.index(name) for name in header}
            region_column = "省份" if "省份" in columns else "地区"
            current_region = ""
            source = source_metadata(project_root, source_path, registry)
            for row in rows[1:]:
                region = row[columns[region_column]]
                current_region = region or current_region
                sequence = row[columns["序号"]]
                if not re.fullmatch(r"\d+", sequence):
                    continue
                entity = row[columns["孵化器名称"]]
                operator = row[columns["运营主体名称"]]
                detail = row[columns["类型"]]
                records.append(
                    base_record(
                        source_id=f"most_incubator_{year}",
                        list_year=year,
                        data_period=str(year),
                        variable="most_new_national_incubator_count",
                        list_name=f"{year}年度国家级科技企业孵化器名单",
                        list_status="confirmed_flow",
                        category="entrepreneurial_ecosystem",
                        province_raw=current_region,
                        location_raw=current_region,
                        entity_name=entity,
                        item_name=operator,
                        detail_type=detail,
                        aggregation_key=entity,
                        source=source,
                        extraction_method="soffice_doc_to_docx_python_docx_table",
                        mapping_text=" ".join([current_region, entity, operator]),
                    )
                )
    return records


def extract_intangible_units(
    project_root: Path,
    converter: LegacyConverter,
    registry: dict[str, dict[str, str]],
) -> list[dict[str, Any]]:
    records = []
    directory = project_root / "03_external_raw" / "mct_intangible_units"
    for source_path in sorted(directory.glob("*.xls")):
        converted = converter.convert(source_path, "xlsx")
        rows = list(iter_xlsx_rows(converted, max_col=7))
        title = " ".join(cell for row in rows[:3] for cell in row if cell)
        if "评估不合格" in title:
            status = "excluded_failed_evaluation"
        elif "评估合格" in title:
            status = "confirmed_active"
        elif "重新认定" in title:
            status = "confirmed_active_rerecognized"
        else:
            status = "[NEEDS REVIEW]"
        header_index = next((i for i, row in enumerate(rows) if "项目名称" in row and "项目编号" in row), None)
        if header_index is None:
            continue
        header = rows[header_index]
        columns = {name: header.index(name) for name in header if name}
        unit_column = next((name for name in columns if "保护单位" in name), "")
        source = source_metadata(project_root, source_path, registry)
        for row in rows[header_index + 1 :]:
            if not row or not re.fullmatch(r"\d+", row[columns.get("序号", 0)]):
                continue
            item = row[columns["项目名称"]]
            item_code = row[columns["项目编号"]]
            item_type = row[columns.get("项目类别", 3)]
            location = row[columns["申报地区或单位"]]
            entity = row[columns[unit_column]] if unit_column else ""
            records.append(
                base_record(
                    source_id="mct_intangible_units",
                    list_year=2023,
                    data_period="2023 current protection-unit evaluation",
                    variable="mct_active_intangible_protection_unit_count",
                    list_name="国家级非物质文化遗产代表性项目保护单位名单",
                    list_status=status,
                    category="cultural_history",
                    province_raw=location,
                    location_raw=location,
                    entity_name=entity,
                    item_name=f"{item} [{item_code}]",
                    detail_type=item_type,
                    aggregation_key=entity,
                    source=source,
                    extraction_method="soffice_xls_to_xlsx_openpyxl",
                    mapping_text=" ".join([location, entity, item]),
                )
            )
    return records


def extract_intangible_productive(
    project_root: Path,
    converter: LegacyConverter,
    registry: dict[str, dict[str, str]],
) -> list[dict[str, Any]]:
    records = []
    directory = project_root / "03_external_raw" / "mct_intangible_productive"
    for source_path in sorted(directory.glob("*.doc")):
        converted = converter.convert(source_path, "docx")
        rows = docx_table(converted)
        if not rows or "推荐企业名称" not in rows[0]:
            continue
        header = rows[0]
        columns = {name: header.index(name) for name in header}
        source = source_metadata(project_root, source_path, registry)
        for row in rows[1:]:
            if not re.fullmatch(r"\d+", row[columns["序号"]]):
                continue
            region = row[columns["推荐地区"]]
            entity = row[columns["推荐企业名称"]]
            item = row[columns["依托国家级非物质文化遗产代表性项目名称"]]
            records.append(
                base_record(
                    source_id="mct_intangible_productive",
                    list_year=2024,
                    data_period="2023-2025",
                    variable="mct_intangible_productive_base_count",
                    list_name="2023—2025年国家级非物质文化遗产生产性保护示范基地名单",
                    list_status="confirmed_period_list",
                    category="cultural_history",
                    province_raw=region,
                    location_raw=region,
                    entity_name=entity,
                    item_name=item,
                    detail_type="",
                    aggregation_key=entity,
                    source=source,
                    extraction_method="soffice_doc_to_docx_python_docx_table",
                    mapping_text=" ".join([region, entity, item]),
                )
            )
    return records


def extract_cultural_industry(
    project_root: Path,
    registry: dict[str, dict[str, str]],
) -> list[dict[str, Any]]:
    records = []
    directory = project_root / "03_external_raw" / "mct_cultural_industry"
    for path in sorted(directory.glob("*.pdf")):
        if "623773137192" in path.name:
            status = "proposal_new_not_final"
            variable = "mct_cultural_industry_proposed_new_base_count"
            list_name = "新一批国家文化产业示范基地拟命名名单"
        elif "24519cdc958b" in path.name:
            status = "proposal_retained_not_final"
            variable = "mct_cultural_industry_proposed_retained_base_count"
            list_name = "拟保留命名国家文化产业示范基地名单"
        else:
            status = "[NEEDS REVIEW]"
            variable = ""
            list_name = "[NEEDS REVIEW]"
        source = source_metadata(project_root, path, registry)
        current_region = ""
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        cells = [clean_text(cell) for cell in row]
                        if len(cells) < 3 or not re.fullmatch(r"\d+", cells[0]):
                            continue
                        current_region = cells[1] or current_region
                        entity = cells[2]
                        records.append(
                            base_record(
                                source_id="mct_cultural_industry",
                                list_year=2024,
                                data_period="2024 proposal notice",
                                variable=variable,
                                list_name=list_name,
                                list_status=status,
                                category="cultural_history",
                                province_raw=current_region,
                                location_raw=current_region,
                                entity_name=entity,
                                item_name="",
                                detail_type="",
                                aggregation_key=entity,
                                source=source,
                                extraction_method="pdfplumber_table_with_merged_cell_carry_forward",
                                mapping_text=" ".join([current_region, entity]),
                            )
                        )
    return records


def extract_miit_culture_technology(
    project_root: Path,
    converter: LegacyConverter,
    registry: dict[str, dict[str, str]],
) -> list[dict[str, Any]]:
    records = []
    directory = project_root / "03_external_raw" / "miit_culture_technology"
    for source_path in sorted(directory.glob("*.doc")):
        converted = converter.convert(source_path, "docx")
        document = Document(converted)
        source = source_metadata(project_root, source_path, registry)
        current_type = ""
        for paragraph in document.paragraphs:
            text = clean_text(paragraph.text)
            if text in {"一、集聚类", "二、单体类"}:
                current_type = text.split("、", 1)[1]
                continue
            match = re.match(r"^(\d+)[.．、]\s*(.+)$", text)
            if not match:
                continue
            entity = match.group(2)
            records.append(
                base_record(
                    source_id="miit_culture_technology",
                    list_year=2024,
                    data_period="2024 fifth batch",
                    variable="miit_culture_technology_fusion_base_count",
                    list_name="第五批国家文化和科技融合示范基地名单",
                    list_status="confirmed_flow",
                    category="cultural_history",
                    province_raw="",
                    location_raw="",
                    entity_name=entity,
                    item_name="",
                    detail_type=current_type,
                    aggregation_key=entity,
                    source=source,
                    extraction_method="soffice_doc_to_docx_python_docx_paragraph_regex",
                    mapping_text=entity,
                )
            )
    return records


def apply_geography(records: list[dict[str, Any]], matcher: DivisionMatcher) -> None:
    for record in records:
        result = matcher.match(
            record["province_raw"],
            record["location_raw"],
            record["mapping_text"],
        )
        record.update(
            {
                "prefecture_code": result.prefecture_code,
                "prefecture_name": result.prefecture_name,
                "prefecture_type": result.prefecture_type,
                "province_code": result.province_code,
                "province_name": result.province_name,
                "match_status": result.status,
                "match_method": result.method,
                "match_evidence": result.evidence,
            }
        )
        flags = {flag for flag in result.flags.split("|") if flag}
        excluded = record["list_status"].startswith("excluded_")
        if excluded:
            flags.add("not_active_list_member")
        if record["list_status"].startswith("proposal_"):
            flags.add("proposal_not_final")
        if record["variable"] == "moe_adult_higher_school_count":
            flags.add("adult_school_location_not_explicit_in_source_column")
        matched = result.status in {"matched_exact", "matched_flagged"}
        record["usable_for_city_panel"] = int(matched and not excluded)
        record["index_eligible"] = int(
            bool(VARIABLE_SPECS.get(record["variable"], {}).get("index_eligible"))
            and matched
            and not excluded
        )
        record["qc_flags"] = "|".join(sorted(flags))


def province_code_for_record(record: dict[str, Any], matcher: DivisionMatcher) -> str:
    if record.get("province_code"):
        return str(record["province_code"])
    code, _ = matcher.find_province(record.get("province_raw", ""))
    return code


def build_coverage(
    records: list[dict[str, Any]],
    matcher: DivisionMatcher,
) -> list[dict[str, Any]]:
    coverage = []
    groups: dict[tuple[str, int], list[dict[str, Any]]] = defaultdict(list)
    for record in records:
        if record["variable"] and not record["list_status"].startswith("excluded_"):
            groups[(record["variable"], int(record["list_year"]))].append(record)
    for (variable, year), group in sorted(groups.items()):
        matched = [row for row in group if row["usable_for_city_panel"]]
        outside_scope = [
            row for row in group if row["match_status"] == "outside_prefecture_universe"
        ]
        unmatched = [
            row
            for row in group
            if not row["usable_for_city_panel"]
            and row["match_status"] != "outside_prefecture_universe"
        ]
        unmatched_unknown_province = sum(
            1 for row in unmatched if not province_code_for_record(row, matcher)
        )
        coverage.append(
            {
                "variable": variable,
                "list_year": year,
                "total_source_records": len(group),
                "matched_records": len(matched),
                "unmatched_records": len(unmatched),
                "outside_prefecture_scope_records": len(outside_scope),
                "match_rate_pct": (
                    round(100 * len(matched) / (len(matched) + len(unmatched)), 3)
                    if matched or unmatched
                    else ""
                ),
                "matched_prefectures": len({row["prefecture_code"] for row in matched}),
                "unmatched_unknown_province": unmatched_unknown_province,
                "safe_zero_rule": (
                    "No zero is filled anywhere"
                    if unmatched_unknown_province
                    else "Zero only where the prefecture's province has no unmatched record"
                ),
                "index_eligible": int(VARIABLE_SPECS[variable]["index_eligible"]),
                "source_status": VARIABLE_SPECS[variable]["status"],
            }
        )
    return coverage


def build_panel(
    records: list[dict[str, Any]],
    universe: list[dict[str, str]],
    matcher: DivisionMatcher,
    start_year: int,
    end_year: int,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    observed_groups: dict[tuple[str, int], list[dict[str, Any]]] = defaultdict(list)
    for record in records:
        if record["variable"] and not record["list_status"].startswith("excluded_"):
            observed_groups[(record["variable"], int(record["list_year"]))].append(record)

    long_rows = []
    wide_rows: dict[tuple[str, int], dict[str, Any]] = {}
    for city in universe:
        for year in range(start_year, end_year + 1):
            key = (city["prefecture_code"], year)
            wide_rows[key] = {
                **{k: city[k] for k in ("prefecture_code", "prefecture_name", "prefecture_type", "province_code", "province_name")},
                "year": year,
            }

    for variable, spec in VARIABLE_SPECS.items():
        observed_years = sorted(year for var, year in observed_groups if var == variable)
        for city in universe:
            city_code = city["prefecture_code"]
            province_code = city["province_code"]
            for year in range(start_year, end_year + 1):
                group = observed_groups.get((variable, year), [])
                value: Any = ""
                if not group:
                    status = "not_observed_year"
                else:
                    matched_city = [
                        row
                        for row in group
                        if row["usable_for_city_panel"] and row["prefecture_code"] == city_code
                    ]
                    aggregation_keys = {row["aggregation_key"] for row in matched_city if row["aggregation_key"]}
                    if aggregation_keys:
                        value = len(aggregation_keys)
                        unmatched_same_province = any(
                            not row["usable_for_city_panel"]
                            and row["match_status"] != "outside_prefecture_universe"
                            and province_code_for_record(row, matcher) == province_code
                            for row in group
                        )
                        status = (
                            "observed_partial_positive" if unmatched_same_province else "observed_matched_positive"
                        )
                    else:
                        unknown_province = any(
                            not row["usable_for_city_panel"]
                            and row["match_status"] != "outside_prefecture_universe"
                            and not province_code_for_record(row, matcher)
                            for row in group
                        )
                        unmatched_same_province = any(
                            not row["usable_for_city_panel"]
                            and row["match_status"] != "outside_prefecture_universe"
                            and province_code_for_record(row, matcher) == province_code
                            for row in group
                        )
                        if unknown_province or unmatched_same_province:
                            value = ""
                            status = "observed_unsafe_zero_blank"
                        else:
                            value = 0
                            status = "observed_safe_zero"
                long_rows.append(
                    {
                        **{k: city[k] for k in ("prefecture_code", "prefecture_name", "prefecture_type", "province_code", "province_name")},
                        "year": year,
                        "variable": variable,
                        "value": value,
                        "coverage_status": status,
                        "source_observed_year": int(year in observed_years),
                        "index_eligible": int(spec["index_eligible"]),
                        "source_status": spec["status"],
                    }
                )
                wide_rows[(city_code, year)][variable] = value
    return list(wide_rows.values()), long_rows


def build_source_inventory(
    project_root: Path,
    registry: dict[str, dict[str, str]],
    records: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    counts = Counter(row["source_id"] for row in records)
    source_ids = [
        "moe_universities",
        "most_incubator_2015",
        "most_incubator_2016",
        "most_incubator_2017",
        "most_incubator_2020",
        "most_incubator_2022",
        "mct_intangible_units",
        "mct_intangible_productive",
        "mct_cultural_industry",
        "miit_culture_technology",
        "miit_future_industry",
    ]
    all_events = []
    events_path = project_root / "01_source_register" / "download_events.csv"
    with events_path.open(encoding="utf-8-sig", newline="") as handle:
        all_events = list(csv.DictReader(handle))
    output = []
    for source_id in source_ids:
        events = [row for row in all_events if row.get("source_id") == source_id]
        errors = [row for row in events if row.get("error")]
        if source_id == "miit_future_industry":
            status = "excluded_no_city_incidence_national_task_topics_only"
            note = (
                "Downloaded DOC is a national clean-hydrogen task specification, not a winner/location list. "
                "Two other task attachments returned HTTP 403; none supplies city incidence."
            )
        else:
            status = "records_extracted" if counts[source_id] else "no_records_extracted"
            note = ""
        output.append(
            {
                "source_id": source_id,
                "download_events": len(events),
                "download_errors": len(errors),
                "extracted_records": counts[source_id],
                "pipeline_status": status,
                "note": note,
            }
        )
    return output


def build_codebook(
    panel_path: Path,
    panel_rows: list[dict[str, Any]],
    out_dir: Path,
    mca_meta: dict[str, str],
) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)
    base_columns = {
        "prefecture_code": ("id", "MCA six-digit prefecture code; municipality uses its provincial code", None),
        "prefecture_name": ("categorical", "Official MCA prefecture name", None),
        "prefecture_type": ("categorical", "Administrative-unit type", None),
        "province_code": ("categorical", "MCA six-digit province code", None),
        "province_name": ("categorical", "Official MCA province name", None),
        "year": ("date", "Calendar/list year", "year"),
    }
    columns = []
    for name, (role, label, units) in base_columns.items():
        values = [row.get(name, "") for row in panel_rows]
        columns.append(
            {
                "name": name,
                "role": role,
                "dtype": "string" if name != "year" else "integer",
                "n": len(values),
                "n_missing": sum(value == "" for value in values),
                "pct_missing": round(100 * sum(value == "" for value in values) / len(values), 3),
                "n_unique": len(set(values)),
                "label": label,
                "units": units,
                "needs_dictionary": False,
                "notes": ["Role is a deterministic draft and requires researcher confirmation."],
            }
        )
    for name, spec in VARIABLE_SPECS.items():
        values = [row.get(name, "") for row in panel_rows]
        numeric = [int(value) for value in values if value != ""]
        columns.append(
            {
                "name": name,
                "role": "continuous",
                "dtype": "nullable_integer",
                "n": len(values),
                "n_missing": sum(value == "" for value in values),
                "pct_missing": round(100 * sum(value == "" for value in values) / len(values), 3),
                "n_unique": len(set(numeric)),
                "label": spec["label"],
                "units": "official listed entities",
                "needs_dictionary": False,
                "stats": {
                    "min": min(numeric) if numeric else None,
                    "max": max(numeric) if numeric else None,
                },
                "notes": [
                    f"Construct: {spec['construct']}.",
                    f"Source status: {spec['status']}.",
                    f"Index eligible: {spec['index_eligible']}.",
                    "Blank is missing/not safely zero, not zero.",
                    "Role assignment is DRAFT_REQUIRES_RESEARCHER_CONFIRMATION.",
                ],
            }
        )
    codebook = {
        "schema_version": 1,
        "status": "DRAFT_REQUIRES_RESEARCHER_ROLE_CONFIRMATION",
        "source": panel_path.as_posix(),
        "n_rows": len(panel_rows),
        "n_columns": len(columns),
        "needs_dictionary_count": 0,
        "geography_source": mca_meta,
        "columns": columns,
    }
    json_path = out_dir / "official_lists_codebook.json"
    json_path.write_text(json.dumps(codebook, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        "# Official Lists Panel Codebook (Draft)",
        "",
        "> Status: `DRAFT_REQUIRES_RESEARCHER_ROLE_CONFIRMATION`. No coded meanings were guessed.",
        "",
        f"- Source: `{panel_path}`",
        f"- Rows: {len(panel_rows):,}",
        f"- Geography source: {mca_meta['source_url']}",
        f"- Geography snapshot SHA-256: `{mca_meta['snapshot_sha256']}`",
        "- Blank count cells mean not observed or not safely zero; they are never implicit zeros.",
        "",
        "| Variable | Role | Label | Units | Missing N | Needs dictionary |",
        "|---|---|---|---|---:|---|",
    ]
    for column in columns:
        lines.append(
            f"| `{column['name']}` | {column['role']} | {column['label']} | "
            f"{column.get('units') or ''} | {column['n_missing']} | no |"
        )
    lines.extend(
        [
            "",
            "## Operationalization cautions",
            "",
            "- University variables are institutional stocks, not research output or technology hard power.",
            "- Incubator and culture/technology-base variables are announcement flows unless explicitly labelled stock.",
            "- Cultural-industry proposal variables are not final awards and are excluded from the core index.",
            "- No stock is forward-filled and no missing cell is imputed.",
            "- Current MCA codes are used for geography. Historical names that cannot be verified remain unmatched.",
        ]
    )
    (out_dir / "official_lists_codebook.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def validate_outputs(
    records: list[dict[str, Any]],
    universe: list[dict[str, str]],
    wide_panel: list[dict[str, Any]],
    long_panel: list[dict[str, Any]],
    start_year: int,
    end_year: int,
) -> dict[str, Any]:
    expected_wide = 337 * (end_year - start_year + 1)
    expected_long = expected_wide * len(VARIABLE_SPECS)
    checks = {
        "prefecture_universe_is_337": len(universe) == 337,
        "wide_panel_complete_city_year_grid": len(wide_panel) == expected_wide,
        "long_panel_complete_city_year_variable_grid": len(long_panel) == expected_long,
        "wide_panel_key_unique": len({(r["prefecture_code"], r["year"]) for r in wide_panel}) == len(wide_panel),
        "air_quality_excluded": all(not r["source_id"].startswith("mee_") for r in records),
        "no_30_city_filter": len({r["prefecture_code"] for r in wide_panel}) == 337,
        "no_unmatched_record_used": all(
            not r["usable_for_city_panel"] or r["match_status"] in {"matched_exact", "matched_flagged"}
            for r in records
        ),
        "failed_intangible_unit_excluded": all(
            not r["usable_for_city_panel"]
            for r in records
            if r["list_status"] == "excluded_failed_evaluation"
        ),
    }
    if not all(checks.values()):
        failed = [name for name, passed in checks.items() if not passed]
        raise AssertionError(f"Official-list validation failed: {failed}")
    return {
        "status": "PASS",
        "checks": checks,
        "expected_wide_rows": expected_wide,
        "expected_long_rows": expected_long,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--project-root", type=Path, default=Path(__file__).resolve().parents[2])
    parser.add_argument("--start-year", type=int, default=2014)
    parser.add_argument("--end-year", type=int, default=2026)
    parser.add_argument("--refresh-mca", action="store_true", help="Refresh the cached MCA HTML snapshot")
    args = parser.parse_args()
    project_root = args.project_root.resolve()
    if args.end_year < args.start_year:
        raise SystemExit("--end-year must be >= --start-year")

    intermediate = project_root / "05_intermediate"
    panel_dir = project_root / "06_panel"
    qc_dir = project_root / "10_qc"
    dictionary_dir = project_root / "04_crosswalk" / "variable_dictionary"
    for directory in (intermediate, panel_dir, qc_dir, dictionary_dir):
        directory.mkdir(parents=True, exist_ok=True)

    snapshot = intermediate / "official_lists_mca_map_snapshot.html"
    page_text, mca_meta = fetch_mca_snapshot(snapshot, args.refresh_mca)
    divisions = parse_mca_divisions(page_text)
    universe = build_prefecture_universe(divisions)
    matcher = DivisionMatcher(divisions, universe)
    registry = read_download_registry(project_root)
    converter = LegacyConverter(intermediate / "official_lists_conversion_cache")

    records: list[dict[str, Any]] = []
    records.extend(extract_moe(project_root, converter, registry))
    records.extend(extract_incubator_html(project_root, registry))
    records.extend(extract_incubator_doc(project_root, converter, registry))
    records.extend(extract_intangible_units(project_root, converter, registry))
    records.extend(extract_intangible_productive(project_root, converter, registry))
    records.extend(extract_cultural_industry(project_root, registry))
    records.extend(extract_miit_culture_technology(project_root, converter, registry))
    apply_geography(records, matcher)
    records.sort(key=lambda r: (r["source_id"], int(r["list_year"]), r["record_id"]))

    universe_path = intermediate / "official_lists_prefecture_universe.csv"
    write_csv(
        universe_path,
        universe,
        [
            "prefecture_code",
            "prefecture_name",
            "prefecture_type",
            "province_code",
            "province_name",
            "pinyin",
            "universe_scope",
        ],
    )
    records_path = intermediate / "official_lists_records.csv"
    write_csv(records_path, records, RECORD_FIELDS)

    unmatched = [row for row in records if not row["usable_for_city_panel"]]
    write_csv(qc_dir / "official_lists_unmatched.csv", unmatched, RECORD_FIELDS)
    coverage = build_coverage(records, matcher)
    coverage_fields = list(coverage[0]) if coverage else []
    write_csv(qc_dir / "official_lists_coverage.csv", coverage, coverage_fields)

    inventory = build_source_inventory(project_root, registry, records)
    write_csv(qc_dir / "official_lists_source_inventory.csv", inventory, list(inventory[0]))

    wide_panel, long_panel = build_panel(
        records, universe, matcher, args.start_year, args.end_year
    )
    wide_fields = [
        "prefecture_code",
        "prefecture_name",
        "prefecture_type",
        "province_code",
        "province_name",
        "year",
        *VARIABLE_SPECS,
    ]
    long_fields = [
        "prefecture_code",
        "prefecture_name",
        "prefecture_type",
        "province_code",
        "province_name",
        "year",
        "variable",
        "value",
        "coverage_status",
        "source_observed_year",
        "index_eligible",
        "source_status",
    ]
    panel_path = panel_dir / "official_lists_city_year.csv"
    write_csv(panel_path, wide_panel, wide_fields)
    write_csv(panel_dir / "official_lists_city_year_long.csv", long_panel, long_fields)

    validation = validate_outputs(
        records, universe, wide_panel, long_panel, args.start_year, args.end_year
    )
    summary = {
        "generated_at_utc": dt.datetime.now(dt.timezone.utc).isoformat(),
        "status": "PASS",
        "geography": {
            "universe_rows": len(universe),
            "type_counts": dict(Counter(row["prefecture_type"] for row in universe)),
            **mca_meta,
        },
        "records": {
            "total": len(records),
            "usable_for_city_panel": sum(int(row["usable_for_city_panel"]) for row in records),
            "index_eligible_records": sum(int(row["index_eligible"]) for row in records),
            "match_status_counts": dict(Counter(row["match_status"] for row in records)),
            "source_counts": dict(Counter(row["source_id"] for row in records)),
        },
        "panel": {
            "start_year": args.start_year,
            "end_year": args.end_year,
            "wide_rows": len(wide_panel),
            "long_rows": len(long_panel),
            "variables": len(VARIABLE_SPECS),
            "zero_policy": "Only province-safe zeros; unsafe zeros and unobserved years are blank.",
        },
        "validation": validation,
    }
    (qc_dir / "official_lists_summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    build_codebook(panel_path, wide_panel, dictionary_dir, mca_meta)

    report_lines = [
        "# Official Lists Pipeline Validation",
        "",
        f"- Status: **{summary['status']}**",
        f"- Geography: {len(universe)} mainland prefecture-equivalent units; no 30-city filter.",
        f"- Years: {args.start_year}–{args.end_year}; wide panel rows: {len(wide_panel):,}.",
        f"- Extracted records: {len(records):,}; city-usable: {summary['records']['usable_for_city_panel']:,}.",
        "- Air quality: explicitly excluded.",
        "- Missing/unsafe zeros: blank, never coerced to zero.",
        "- Geography vintage: live MCA snapshot; source page does not expose an effective date.",
        "",
        "## Source limitations",
        "",
        "- MIIT future-industry attachment is a national task specification without city incidence and is excluded.",
        "- Cultural-industry PDFs are proposal-stage public notices; their variables are not core-index eligible.",
        "- Adult-school, productive-protection-base, and several incubator records lack explicit city columns; unmatched records remain in the QC file.",
        "- Current MCA geography does not silently rewrite unverifiable historical place names.",
        "",
        "## Skill gates",
        "",
        "- `clean-data`: only profiling/flagging was applied; no user-unapproved cleaning, deletion, or imputation.",
        "- `generate-codebook`: codebook is a draft and must receive researcher confirmation before being authoritative.",
        "- `define-variables`: proposal-stage and structurally non-city sources are not promoted into the core index.",
    ]
    (qc_dir / "official_lists_validation.md").write_text(
        "\n".join(report_lines) + "\n", encoding="utf-8"
    )

    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
