#!/usr/bin/env python3
"""Build an auditable text/city-mention index for every collected snapshot.

This is a staging layer, not a finished index variable.  Exact official city
names are marked as high-confidence mentions; suffix-free names are retained as
review candidates only.  That distinction prevents generic web text from
silently becoming a city-level measurement.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import html
import json
import re
import shutil
import subprocess
import tempfile
import zipfile
from collections import Counter
from datetime import datetime, timezone
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Iterable


PROJECT = Path(__file__).resolve().parents[2]
EVENTS = PROJECT / "01_source_register" / "download_events.csv"
SEEDS = PROJECT / "01_source_register" / "source_seeds.csv"
CITY_MASTER = PROJECT / "04_crosswalk" / "city_master_297_snapshot.csv"
DOCUMENTS = PROJECT / "05_intermediate" / "full_collection_document_registry.csv"
MENTIONS = PROJECT / "05_intermediate" / "full_collection_city_mentions.csv"
COVERAGE = PROJECT / "10_qc" / "full_collection_city_mention_coverage.csv"
SUMMARY = PROJECT / "10_qc" / "full_collection_text_index_summary.json"


DOCUMENT_FIELDS = [
    "document_id", "source_id", "category", "agency", "source_url",
    "source_file", "source_sha256", "content_type", "bytes",
    "fetched_at_utc", "extraction_status", "extraction_method",
    "text_characters", "detected_years", "exact_city_count",
    "candidate_city_count", "error",
]

MENTION_FIELDS = [
    "mention_id", "document_id", "source_id", "category", "source_file",
    "source_url", "source_sha256", "city_code", "city_name",
    "province_code", "province_name", "match_method", "mention_text",
    "mention_count", "detected_years", "formal_variable_eligible",
]

COVERAGE_FIELDS = [
    "source_id", "category", "documents", "documents_text_extracted",
    "documents_failed", "exact_cities", "candidate_cities", "exact_mentions",
    "candidate_mentions",
]


class VisibleHTML(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.hidden = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.lower() in {"script", "style", "noscript", "svg"}:
            self.hidden += 1

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in {"script", "style", "noscript", "svg"} and self.hidden:
            self.hidden -= 1

    def handle_data(self, data: str) -> None:
        if not self.hidden:
            self.parts.append(data)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def write_csv(path: Path, rows: Iterable[dict[str, Any]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def stable_id(*parts: str) -> str:
    return hashlib.sha256("\x1f".join(parts).encode("utf-8")).hexdigest()[:24]


def decode_bytes(payload: bytes, content_type: str = "") -> str:
    encodings: list[str] = []
    match = re.search(r"charset=([\w-]+)", content_type, re.I)
    if match:
        encodings.append(match.group(1))
    head = payload[:8192].decode("ascii", errors="ignore")
    match = re.search(r"charset=[\"']?([\w-]+)", head, re.I)
    if match:
        encodings.append(match.group(1))
    encodings.extend(["utf-8", "gb18030"])
    for encoding in encodings:
        try:
            return payload.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return payload.decode("utf-8", errors="replace")


def normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", html.unescape(value)).strip()


def extract_html(path: Path, content_type: str) -> tuple[str, str]:
    parser = VisibleHTML()
    parser.feed(decode_bytes(path.read_bytes(), content_type))
    return normalize_text(" ".join(parser.parts)), "stdlib_html_parser"


def extract_pdf(path: Path) -> tuple[str, str]:
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        return normalize_text(" ".join((page.extract_text() or "") for page in reader.pages)), "pypdf"
    except ImportError:
        command = shutil.which("pdftotext")
        if not command:
            raise RuntimeError("missing pypdf and pdftotext")
        completed = subprocess.run(
            [command, "-layout", str(path), "-"], capture_output=True, check=True
        )
        return normalize_text(decode_bytes(completed.stdout)), "pdftotext"


def extract_docx(path: Path) -> tuple[str, str]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("missing python-docx") from exc
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return normalize_text(" ".join(parts)), "python_docx"


def extract_xlsx(path: Path) -> tuple[str, str]:
    try:
        import openpyxl
    except ImportError as exc:
        raise RuntimeError("missing openpyxl") from exc
    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(sheet.title)
        for row in sheet.iter_rows(values_only=True):
            parts.append(" | ".join("" if value is None else str(value) for value in row))
    workbook.close()
    return normalize_text(" ".join(parts)), "openpyxl"


def extract_xls(path: Path) -> tuple[str, str]:
    try:
        import xlrd
    except ImportError as exc:
        raise RuntimeError("missing xlrd") from exc
    workbook = xlrd.open_workbook(path, on_demand=True)
    parts: list[str] = []
    for sheet in workbook.sheets():
        parts.append(sheet.name)
        for row_idx in range(sheet.nrows):
            parts.append(" | ".join(str(value) for value in sheet.row_values(row_idx)))
    workbook.release_resources()
    return normalize_text(" ".join(parts)), "xlrd"


def extract_doc(path: Path) -> tuple[str, str]:
    command = shutil.which("antiword")
    if command:
        completed = subprocess.run([command, str(path)], capture_output=True, check=False)
        if completed.returncode == 0 and completed.stdout.strip():
            return normalize_text(decode_bytes(completed.stdout)), "antiword"
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("legacy .doc is incompatible with antiword and LibreOffice is missing")
    with tempfile.TemporaryDirectory(prefix="full_collection_doc_") as tmp:
        tmp_root = Path(tmp)
        profile = (tmp_root / "profile").as_uri()
        completed = subprocess.run(
            [
                soffice,
                "--headless",
                f"-env:UserInstallation={profile}",
                "--convert-to",
                "docx",
                "--outdir",
                str(tmp_root),
                str(path),
            ],
            capture_output=True,
            check=False,
        )
        converted = list(tmp_root.glob("*.docx"))
        if completed.returncode != 0 or len(converted) != 1:
            raise RuntimeError(
                "legacy .doc conversion failed: "
                + decode_bytes(completed.stderr or completed.stdout)[-500:]
            )
        text, _ = extract_docx(converted[0])
        return text, "libreoffice_doc_to_docx+python_docx"


def extract_zip(path: Path) -> tuple[str, str]:
    parts: list[str] = []
    methods: set[str] = set()
    with zipfile.ZipFile(path) as archive, tempfile.TemporaryDirectory(prefix="full_collection_zip_") as tmp:
        tmp_root = Path(tmp)
        members = [member for member in archive.infolist() if not member.is_dir()]
        if sum(member.file_size for member in members) > 500 * 1024 * 1024:
            raise RuntimeError("archive expands beyond the 500 MB safety limit")
        for index, member in enumerate(members):
            suffix = Path(member.filename).suffix.lower()
            if suffix not in {".html", ".htm", ".shtml", ".pdf", ".doc", ".docx", ".xls", ".xlsx", ".txt", ".csv"}:
                continue
            target = tmp_root / f"{index}{suffix}"
            target.write_bytes(archive.read(member))
            try:
                member_text, method = extract_text(target, "")
            except Exception:
                continue
            parts.append(member.filename)
            parts.append(member_text)
            methods.add(method)
    if not parts:
        raise RuntimeError("archive contains no supported readable members")
    return normalize_text(" ".join(parts)), "zip_members:" + "+".join(sorted(methods))


def extract_text(path: Path, content_type: str) -> tuple[str, str]:
    suffix = path.suffix.lower()
    lower_type = content_type.lower()
    if "html" in lower_type or suffix in {".html", ".htm", ".shtml"}:
        return extract_html(path, content_type)
    if "pdf" in lower_type or suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx" or "wordprocessingml" in lower_type:
        return extract_docx(path)
    if suffix == ".doc" or "msword" in lower_type:
        return extract_doc(path)
    if suffix == ".xlsx" or "spreadsheetml" in lower_type:
        return extract_xlsx(path)
    if suffix == ".xls" or "ms-excel" in lower_type or "vnd.ms-excel" in lower_type:
        return extract_xls(path)
    if suffix in {".zip"} or lower_type in {"application/zip", "application/x-zip-compressed"}:
        return extract_zip(path)
    if suffix in {".txt", ".csv", ".json"} or lower_type.startswith("text/") or "json" in lower_type:
        return normalize_text(decode_bytes(path.read_bytes(), content_type)), "plain_text_decoder"
    raise RuntimeError(f"unsupported format: {suffix or content_type}")


def city_mentions(text: str, cities: list[dict[str, str]]) -> list[dict[str, Any]]:
    output: list[dict[str, Any]] = []
    for city in cities:
        full = city["city_name"]
        full_count = text.count(full)
        if full_count:
            output.append({
                **city,
                "match_method": "exact_official_city_name",
                "mention_text": full,
                "mention_count": full_count,
                "formal_variable_eligible": 0,
            })
            continue
        alias = re.sub(r"市$", "", full)
        if len(alias) >= 2:
            alias_count = text.count(alias)
            if alias_count:
                output.append({
                    **city,
                    "match_method": "suffix_free_candidate",
                    "mention_text": alias,
                    "mention_count": alias_count,
                    "formal_variable_eligible": 0,
                })
    return output


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--exclude-source", action="append", default=[])
    args = parser.parse_args()
    excluded = set(args.exclude_source)

    seeds = {row["source_id"]: row for row in read_csv(SEEDS)}
    cities = read_csv(CITY_MASTER)
    successful: dict[str, dict[str, str]] = {}
    for event in read_csv(EVENTS):
        if (
            event.get("source_id") not in excluded
            and event.get("http_status", "").startswith("2")
            and event.get("saved_path")
        ):
            successful[event["saved_path"]] = event

    documents: list[dict[str, Any]] = []
    mentions: list[dict[str, Any]] = []
    for saved_path, event in sorted(successful.items()):
        path = PROJECT / saved_path
        source_id = event["source_id"]
        seed = seeds.get(source_id, {})
        document_id = stable_id(source_id, saved_path, event.get("sha256", ""))
        error = ""
        method = ""
        text = ""
        status = "failed"
        if not path.is_file():
            error = "registered file missing"
        else:
            try:
                text, method = extract_text(path, event.get("content_type", ""))
                status = "extracted" if text else "empty_text"
            except Exception as exc:  # file-level failure must not stop the batch
                error = f"{type(exc).__name__}: {exc}"
        years = sorted(set(re.findall(r"(?<!\d)(?:19[5-9]\d|20[0-3]\d)(?!\d)", text)))
        found = city_mentions(text, cities) if text else []
        exact = [row for row in found if row["match_method"] == "exact_official_city_name"]
        candidates = [row for row in found if row["match_method"] == "suffix_free_candidate"]
        documents.append({
            "document_id": document_id,
            "source_id": source_id,
            "category": seed.get("category", ""),
            "agency": seed.get("agency", ""),
            "source_url": event.get("url", ""),
            "source_file": saved_path,
            "source_sha256": event.get("sha256", ""),
            "content_type": event.get("content_type", ""),
            "bytes": event.get("bytes", ""),
            "fetched_at_utc": event.get("fetched_at_utc", ""),
            "extraction_status": status,
            "extraction_method": method,
            "text_characters": len(text),
            "detected_years": "|".join(years),
            "exact_city_count": len(exact),
            "candidate_city_count": len(candidates),
            "error": error,
        })
        for found_row in found:
            mention_id = stable_id(document_id, found_row["city_code"], found_row["match_method"])
            mentions.append({
                "mention_id": mention_id,
                "document_id": document_id,
                "source_id": source_id,
                "category": seed.get("category", ""),
                "source_file": saved_path,
                "source_url": event.get("url", ""),
                "source_sha256": event.get("sha256", ""),
                "city_code": found_row["city_code"],
                "city_name": found_row["city_name"],
                "province_code": found_row["province_code"],
                "province_name": found_row["province_name"],
                "match_method": found_row["match_method"],
                "mention_text": found_row["mention_text"],
                "mention_count": found_row["mention_count"],
                "detected_years": "|".join(years),
                "formal_variable_eligible": found_row["formal_variable_eligible"],
            })

    coverage_rows: list[dict[str, Any]] = []
    source_ids = sorted({row["source_id"] for row in documents})
    for source_id in source_ids:
        source_documents = [row for row in documents if row["source_id"] == source_id]
        source_mentions = [row for row in mentions if row["source_id"] == source_id]
        exact_mentions = [row for row in source_mentions if row["match_method"] == "exact_official_city_name"]
        candidate_mentions = [row for row in source_mentions if row["match_method"] == "suffix_free_candidate"]
        coverage_rows.append({
            "source_id": source_id,
            "category": seeds.get(source_id, {}).get("category", ""),
            "documents": len(source_documents),
            "documents_text_extracted": sum(row["extraction_status"] == "extracted" for row in source_documents),
            "documents_failed": sum(row["extraction_status"] == "failed" for row in source_documents),
            "exact_cities": len({row["city_code"] for row in exact_mentions}),
            "candidate_cities": len({row["city_code"] for row in candidate_mentions}),
            "exact_mentions": sum(int(row["mention_count"]) for row in exact_mentions),
            "candidate_mentions": sum(int(row["mention_count"]) for row in candidate_mentions),
        })

    write_csv(DOCUMENTS, documents, DOCUMENT_FIELDS)
    write_csv(MENTIONS, mentions, MENTION_FIELDS)
    write_csv(COVERAGE, coverage_rows, COVERAGE_FIELDS)
    status_counts = Counter(row["extraction_status"] for row in documents)
    summary = {
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "city_universe": "293 prefecture-level cities plus 4 municipalities",
        "city_universe_rows": len(cities),
        "excluded_sources": sorted(excluded),
        "documents": len(documents),
        "document_status_counts": dict(sorted(status_counts.items())),
        "sources_with_documents": len(source_ids),
        "city_mentions": len(mentions),
        "exact_city_mentions": sum(row["match_method"] == "exact_official_city_name" for row in mentions),
        "suffix_free_candidates": sum(row["match_method"] == "suffix_free_candidate" for row in mentions),
        "formal_variable_eligible_mentions": 0,
        "note": "This is discovery/QC staging. Formal variables require source-specific record extraction and city mapping.",
    }
    SUMMARY.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
